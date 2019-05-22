from docx import Document
import re
import os


# 根据内容特征提取内容，并格式化输出
def format_strings(s):
    answer = "".join(s)
    answer = answer.replace('\n', '').replace('\t', '').replace(' ', '').replace('（填空题）', '').replace('\u3000', '')\
        .replace('。', '').replace('null', '').replace(':', '').replace('：', '').replace('姓名', '').replace('电话', '')\
        .replace('(', '').replace(')', '').replace('\xa0', '').replace('?', '').replace('？', '')

    # answer = answer.strip()
    if len(answer) < 1:
        answer = 'null'
    return answer


def each_file(filename, ifdelete=0):
    document = Document(filename)

    all_answers_dict = {}
    default_value = 0
    for i in range(1, 30):
        if i == 5:
            all_answers_dict['5A'] = default_value
            all_answers_dict['5B'] = default_value
            all_answers_dict['5C'] = default_value
            all_answers_dict['5D'] = default_value
        elif i == 13:
            all_answers_dict['13A'] = default_value
            all_answers_dict['13B'] = default_value
            all_answers_dict['13C'] = default_value
            all_answers_dict['13D'] = default_value
            all_answers_dict['13E'] = default_value
            all_answers_dict['13F'] = default_value
        elif i == 14:
            all_answers_dict['14A'] = default_value
            all_answers_dict['14B'] = default_value
            all_answers_dict['14C'] = default_value
            all_answers_dict['14D'] = default_value
            all_answers_dict['14E'] = default_value
            all_answers_dict['14F'] = default_value
            all_answers_dict['14G'] = default_value
            all_answers_dict['14H'] = default_value
            all_answers_dict['14I'] = default_value
        elif i == 20:
            all_answers_dict['20A'] = default_value
            all_answers_dict['20B'] = default_value
            all_answers_dict['20C'] = default_value
            all_answers_dict['20D'] = default_value
        elif i == 21:
            all_answers_dict['21A'] = default_value
            all_answers_dict['21B'] = default_value
            all_answers_dict['21C'] = default_value
            all_answers_dict['21D'] = default_value
            all_answers_dict['21E'] = default_value
            all_answers_dict['21F'] = default_value
            all_answers_dict['21G'] = default_value
            all_answers_dict['21H'] = default_value
            all_answers_dict['21I'] = default_value
        elif i == 22:
            all_answers_dict['22A'] = default_value
            all_answers_dict['22B'] = default_value
            all_answers_dict['22C'] = default_value
            all_answers_dict['22D'] = default_value
        else:
            all_answers_dict[str(i)] = default_value

    list_lines_all = []
    for paragraph in document.paragraphs:
        content = paragraph.text
        if len(content) > 5:
            list_lines_all.append(content)
    # print(list_lines_all)

    # 22题表格数据的提取
    for table in document.tables:
        if table is None:
            break
        # table 的rows和columns得到这个表格的行数和列数
        if len(table.rows) != 5 and len(table.columns) != 6:
            break
        # 遍历表格
        i = 0
        for row in table.rows:
            if len(str(row.cells[0].text).strip()) < 1:  # 表格首列内容为空忽略
                i = i + 1
                continue
            j = 0
            result = 0
            for cell in row.cells:
                if len(str(cell.text).strip()) > 5:
                    j = j + 1
                    continue
                value = str(cell.text).strip()  # 表格内容
                value = value.replace('\n', '').replace(' ', '').replace('\u3000', '').replace('\t', '')
                if len(value) > 0:
                    result = j
                j = j + 1

            if i == 1 and result > 0:
                all_answers_dict['22A'] = result
            if i == 2 and result > 0:
                all_answers_dict['22B'] = result
            if i == 3 and result > 0:
                all_answers_dict['22C'] = result
            if i == 4 and result > 0:
                all_answers_dict['22D'] = result
            i = i + 1
    '''for content in list_lines_all:
        result = re.search(r'软件厂商总体印象.*', content)
        if result:
            #result = str(result.group()).replace('软件厂商总体印象', '').replace('\t', '').replace(' ', '')
            print(result)
'''
    # 单选题的答案提取
    single_list = [1, 2, 3, 4, 6, 7, 8, 9, 10, 11, 12, 15, 16, 17, 18, 19]
    value_list = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L', 'M', 'N']
    for key in ['13A', '13B', '13C', '13D', '13E', '13F', '14A', '14B', '14C', '14D', '14E', '14F',
                '14G', '14F', '14G', '14H', '14I', '20A', '20B', '20C', '20D', '21A', '21B', '21C', '21D',
                '21E', '21F', '21G', '21H', '21I']:
        all_answers_dict[key] = 0
    for content in list_lines_all:
        p_number = re.compile(r'\d{1,2}')  # 最小匹配1或2位数字
        p_answer = re.compile(r'[(](.*?)[)]|[（](.*?)[）]', re.S)  # 最小匹配括号内的字母
        #p_answer = re.compile(r'[（](\s*[a-zA-Z])[）]|[（](\s*[a-zA-Z]\s*)[）]', re.S)  # 最小匹配括号内的字母
        obj_number = re.match(p_number, content)  # 从开头匹配
        obj_answer = re.search(p_answer, content)  # 整个字符串，返回首个匹配
        if obj_answer and obj_number:
            number = obj_number.group()
            answer = re.sub(r'\W', "", obj_answer.group())  # 移除非字母数据的任何字符
            if int(number) in single_list and len(answer) > 0:
                answer = answer[0:1]   # 解决 “K水电费是的发生”这类字符串的问题
                if all_answers_dict[number] == 0:  # 防止重复赋值的问题，匹配结果不止一个，优先第一个
                    answer = answer.upper()
                    '''if len(answer) > 0:
                        index = value_list.index(answer)
                    else:
                        index = -1
                    if index != -1:
                        answer = index + 1
                    else:
                        answer = 0'''
                    all_answers_dict[number] = answer
            elif int(number) in [13, 14, 20, 21]:
                answer = answer.upper()
                # print(number + answer)
                for a in answer:
                    key = number + a
                    all_answers_dict[key] = 1

    # 第5题填空题答案提取
    for content in list_lines_all:
        p_answer1 = re.compile(r'(A.高级职称.*?人)|(高级职称.*?人)')  # 最小匹配1或2位数字
        p_answer2 = re.compile(r'(B.中级职称.*?人)|(中级职称.*?人)')  # 最小匹配1或2位数字
        p_answer3 = re.compile(r'(C.初级职称.*?人)|(初级职称\s*[0-9].*?人)')  # 最小匹配1或2位数字
        p_answer4 = re.compile(r'(D.初级职称以下.*?人)|(初级职称以下.*?人)')  # 最小匹配1或2位数字

        obj_answer1 = re.findall(p_answer1, content)
        obj_answer2 = re.findall(p_answer2, content)
        obj_answer3 = re.findall(p_answer3, content)
        obj_answer4 = re.findall(p_answer4, content)
        if obj_answer1:
            answer = "".join(obj_answer1[0])
            answer = answer.replace('\xa0', '')
            answer = re.sub(r'\D', "", answer)  # 移除非数据的任何字符
            if len(answer) < 1:
                answer = 0
            else:
                answer = int(answer)
            if all_answers_dict['5A'] == 0:
                all_answers_dict['5A'] = answer
        if obj_answer2:
            answer = "".join(obj_answer2[0])
            answer = answer.replace('\xa0', '')
            answer = re.sub(r'\D', "", answer)  # 移除非数据的任何字符
            if len(answer) < 1:
                answer = 0
            else:
                answer = int(answer)
            if all_answers_dict['5B'] == 0:
                all_answers_dict['5B'] = answer
        if obj_answer3:
            answer = "".join(obj_answer3[0])
            answer = answer.replace('\xa0', '')
            answer = re.sub(r'\D', "", answer)  # 移除非数据的任何字符
            if len(answer) < 1:
                answer = 0
            else:
                answer = int(answer)
            if all_answers_dict['5C'] == 0:
                all_answers_dict['5C'] = answer
        if obj_answer4:
            answer = "".join(obj_answer4[0])
            answer = answer.replace('\xa0', '')
            answer = re.sub(r'\D', "", answer)  # 移除非数据的任何字符
            if len(answer) < 1:
                answer = 0
            else:
                answer = int(answer)
            if all_answers_dict['5D'] == 0:
                all_answers_dict['5D'] = answer

    # 23-26填空题答案提取
    contents = list_lines_all
    #print(contents)
    if len(contents) > 26:
        str_23 = '在新旧制度衔接、实施《政府会计制度》过程中，有何心得或成功经验?有何困难和问题?对问题有何建议或解决方案'
        str_24 = '就做好新旧会计制度衔接和《政府会计制度》实施工作，对上级机构和有关主管部门的工作有何意见或建议'
        str_25 = '为便于统计调研任务完成情况，请留下您单位的全称'
        str_26 = '为便于联系，请留下您的姓名以及可联系到您的电话或手机号码'
        str_last = '再次感谢您的参与和配合'
        list_23 = [i for i, x in enumerate(contents) if x.find(str_23) != -1]
        list_24 = [i for i, x in enumerate(contents) if x.find(str_24) != -1]
        list_25 = [i for i, x in enumerate(contents) if x.find(str_25) != -1]
        list_26 = [i for i, x in enumerate(contents) if x.find(str_26) != -1]
        list_last = [i for i, x in enumerate(contents) if x.find(str_last) != -1]
        # 23题
        try:
            if len(list_23) > 0:
                # 答案与题目在同一行
                #print(contents[list_23[0]])
                answer = str(contents[list_23[0]]).replace(str_23, '').replace('（填空题，空间不够填写可另附页）', '').replace('23.', '')
                # 答案与题目不在同一行
                start = int(list_23[0]) + 1
                end = int(list_24[0])
                answer = answer + format_strings(contents[start:end])
                all_answers_dict['23'] = format_strings(answer)
                #print(answer)
            else:
                pass

        except Exception as e:
            pass
        # 24题
        try:
            if len(list_24) > 0:
                # 答案与题目在同一行
                #print(contents[list_24[0]])
                answer = str(contents[list_24[0]]).replace(str_24, '').replace('（填空题，空间不够填写可另附页）', '').replace('24.', '')
                # 答案与题目不在同一行
                start = int(list_24[0]) + 1
                end = int(list_25[0])
                answer = answer + format_strings(contents[start:end])
                all_answers_dict['24'] = format_strings(answer)
                #print(answer)
        except Exception as e:
            pass
        # 25题
        try:
            if len(list_25) > 0:
                # 答案与题目在同一行
                #print(contents[list_25[0]])
                answer = str(contents[list_25[0]]).replace(str_25, '').replace('（填空题）', '').replace('25.', '')
                # 答案与题目不在同一行
                start = int(list_25[0]) + 1
                end = int(list_26[0])
                answer = answer + format_strings(contents[start:end])
                all_answers_dict['25'] = format_strings(answer)
                #print(answer)
                #print(contents[start:end])
        except Exception as e:
            pass
        # 26题
        try:
            # 提取姓名,电话
            if len(list_26) > 0:
                # 答案与题目在同一行
                #print(contents[list_26[0]])
                answer = str(contents[list_26[0]]).replace('为便于联系，请留下您的姓名以及可联系到您的电话或手机', '')\
                    .replace('（填空题）', '').replace('26.', '').replace('\xa0', '')
                    # .replace('联系', '').replace('办公', '').replace('号码', '')\
                    # .replace('手机', '').replace('及', '').replace('联系方式', '').replace('是', '').replace('会计员', '')
                if len(list_last) > 0:
                    # 答案与题目不在同一行
                    start = int(list_26[0]) + 1
                    end = int(list_last[0])
                    s = "".join(contents[start:end])
                    answer = answer + s
                    # answer = format_strings(answer)
                    #print(answer)
                    # 存在多组电话时，应该用符号隔开？？？？？
                    r = re.compile(r'\d{3,4}[-]+\d{6,7}|\d{3,4}[—－]+\d{6,7}|\d{11}|\d{7,10}|[（]\d{3,4}[）]\d{6,7}')
                    phone = re.findall(r, answer)
                    name = re.sub(r, ',', answer)
                    name = name.replace('联系', '').replace('办公', '').replace('号码', '').replace('手机', '').replace('及', '')\
                        .replace('联系方式', '').replace('是', '').replace('会计员', '').replace('再次感谢您的参与和配合！', '').replace('答', '')
                    name = format_strings(name)
                    #print(name)
                    #print("".join(phone))
                    all_answers_dict['26'] = ",".join(phone)
                    all_answers_dict['27'] = name

            '''# 提取电话号码
            n = list_26[0]
            while n < len(contents):
                phone = contents[n]
                phone = format_strings(phone)
                phone = re.findall(r'\d{3,4}[-]+\d{7}|\d{7,12}', phone)
                if phone:
                    all_answers_dict['26'] = "".join(phone)
                    break
                n = n + 1'''
        except Exception as e:
            pass
        sorted(all_answers_dict.keys())
    if ifdelete == 1:
        os.remove(filename)
    print(all_answers_dict)
    return all_answers_dict


if __name__ == '__main__':
    filename = "D:\\datasource\\test\\2016.docx"
    each_file(filename, 0)